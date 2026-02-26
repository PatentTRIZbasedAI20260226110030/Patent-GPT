from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.models.patent_draft import PatentDraft


def export_to_docx(draft: PatentDraft, output_path: str) -> str:
    doc = DocxDocument()

    # Title
    title_para = doc.add_heading(level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"[발명의 명칭] {draft.title}")
    run.font.size = Pt(16)

    # Abstract
    doc.add_heading("요약", level=1)
    doc.add_paragraph(draft.abstract)

    # Background
    doc.add_heading("발명의 배경", level=1)
    doc.add_paragraph(draft.background)

    # Problem Statement
    doc.add_heading("해결하려는 과제", level=1)
    doc.add_paragraph(draft.problem_statement)

    # Solution
    doc.add_heading("과제의 해결 수단", level=1)
    doc.add_paragraph(draft.solution)

    # Claims
    doc.add_heading("청구항", level=1)
    for i, claim in enumerate(draft.claims, 1):
        doc.add_paragraph(f"[항 {i}] {claim}")

    # Effects
    doc.add_heading("발명의 효과", level=1)
    doc.add_paragraph(draft.effects)

    doc.save(output_path)
    return output_path
